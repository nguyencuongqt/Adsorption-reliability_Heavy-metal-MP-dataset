from pathlib import Path

from docx import Document
from docx.enum.text import WD_COLOR_INDEX


ROOT = Path(__file__).resolve().parents[1]
MANUSCRIPT_DIR = ROOT / "manuscript_package" / "Submission_Journal of Environmental Informatics"
MAIN_IN = MANUSCRIPT_DIR / "01_Main_Manuscript_JEI_revised_consistent_final.docx"
SI_IN = MANUSCRIPT_DIR / "02_Supplementary_Information_JEI_revised_consistent.docx"
MAIN_OUT = MANUSCRIPT_DIR / "01_Main_Manuscript_JEI_revised_consistent_final_external_validation.docx"
SI_OUT = MANUSCRIPT_DIR / "02_Supplementary_Information_JEI_revised_consistent_external_validation.docx"


def highlight(paragraph):
    for run in paragraph.runs:
        run.font.highlight_color = WD_COLOR_INDEX.YELLOW


def after_paragraph(document, anchor, text, style_name=None):
    paragraph = document.add_paragraph()
    if style_name:
        paragraph.style = style_name
    paragraph.add_run(text)
    anchor._p.addnext(paragraph._p)
    return paragraph


def after_table(table, new_table):
    table._tbl.addnext(new_table._tbl)


def add_main():
    document = Document(MAIN_IN)
    anchor = next(
        p for p in document.paragraphs
        if p.text.startswith("Finally, the study has implications for data generation and synthesis.")
    )
    text = (
        "As an exploratory external transportability check, the prespecified model pipelines "
        "were refit using all 1,009 locked observations and evaluated in two independent 2026 "
        "Cd(II) studies (35 digitized equilibrium observations). All study-specific R2 values "
        "were negative (Table S10), despite a near-zero pooled R2 for EN (0.043). This analysis "
        "is not a definitive external validation, given only two source studies, Cd-only coverage, "
        "digitized response data, universally unreported surface area, and unreported pH for the "
        "PE study. Nonetheless, it is consistent with the internal grouped-validation evidence "
        "that transfer is constrained by cross-study domain shift and incomplete reporting."
    )
    paragraph = after_paragraph(document, anchor, text, anchor.style.name)
    highlight(paragraph)
    document.save(MAIN_OUT)


def add_si():
    document = Document(SI_IN)
    anchor = next(
        p for p in document.paragraphs
        if p.text.startswith("Finally, the null audit was repeated for three generation seeds")
    )
    heading = after_paragraph(
        document, anchor, "Text S5. Exploratory external transportability check", "Heading 2"
    )
    highlight(heading)
    paragraph_1 = after_paragraph(
        document,
        heading,
        (
            "To complement the internal LOSO analysis, we conducted an exploratory external "
            "transportability check using two independent 2026 Cd(II) adsorption studies not "
            "included in the locked corpus. The external set comprised 35 equilibrium observations: "
            "five UV-C-aged PET observations and 30 unaged/naturally aged PE observations. The "
            "prespecified EN, LGBM, and MLP pipelines were refit once on all 1,009 locked observations, "
            "with the locked 25-predictor representation and 24 documented transformations/interactions "
            "retained; no external observation was used for fitting or model selection. Reported Ce, "
            "pH, temperature, rpm, metal, polymer, aging state, and article-supported functional-group "
            "indicators were mapped to the existing feature contract. Surface area was unreported for "
            "all external observations and pH was unreported for the PE study; these variables were "
            "retained as missing and handled only by the locked pipeline's development-data median-"
            "imputation/missing-indicator procedure. Equilibrium values were digitized from figures, "
            "and the PE Ce values were calculated by mass balance from the article's stated solid-to-"
            "liquid ratio."
        ),
        "Normal",
    )
    highlight(paragraph_1)
    paragraph_2 = after_paragraph(
        document,
        paragraph_1,
        (
            "The results did not demonstrate reliable external transfer (Table S10). Every study-"
            "specific R2 value was negative. Although pooled EN R2 was 0.043 (RMSE = 0.489 mg g-1; "
            "MAE = 0.425 mg g-1), its pooled metric combines two studies with distinct response "
            "distributions and should not be interpreted as evidence of within-study transportability. "
            "We therefore treat this check as exploratory and descriptive, not as a confirmatory "
            "external validation. It instead reinforces the need for independent multi-study test sets "
            "with complete pH and surface-area reporting and broader metal-polymer coverage."
        ),
        "Normal",
    )
    highlight(paragraph_2)

    s9_table = document.tables[12]
    caption = document.add_paragraph(
        "Table S10. Exploratory external transportability check using two independent 2026 Cd(II) studies (35 observations).",
    )
    caption.style = next(p for p in document.paragraphs if p.text.startswith("Table S9.")).style.name
    caption.paragraph_format.page_break_before = True
    s9_table._tbl.addnext(caption._p)
    highlight(caption)

    table = document.add_table(rows=7, cols=4)
    table.style = s9_table.style
    values = [
        ["Metric", "EN", "LGBM", "MLP"],
        ["Pooled R2", "0.043", "-0.058", "-0.331"],
        ["Pooled RMSE (mg g-1)", "0.489", "0.514", "0.576"],
        ["Pooled MAE (mg g-1)", "0.425", "0.388", "0.273"],
        ["Study-specific R2 (PET; n = 5)", "-2.333", "-3.914", "-7.826"],
        ["Study-specific R2 (PE; n = 30)", "-22.613", "-16.040", "-0.910"],
        ["Equal-study mean R2", "-12.473", "-9.977", "-4.368"],
    ]
    for row, row_values in zip(table.rows, values):
        for cell, value in zip(row.cells, row_values):
            cell.text = value
            for paragraph in cell.paragraphs:
                highlight(paragraph)
    caption._p.addnext(table._tbl)
    note = document.add_paragraph(
        (
            "Note: Models were refit on the complete locked development dataset (n = 1,009) using "
            "the prespecified release configurations, without external-data tuning. All external "
            "source-specific R2 values were negative. Pooled metrics combine two studies and are "
            "descriptive. Source data: https://doi.org/10.3390/su18104642 and "
            "https://doi.org/10.1007/s11270-026-09284-9. Surface area was unreported for all external "
            "observations; pH was unreported for the PE study."
        )
    )
    table._tbl.addnext(note._p)
    highlight(note)
    document.save(SI_OUT)


if __name__ == "__main__":
    add_main()
    add_si()
    print(MAIN_OUT)
    print(SI_OUT)
