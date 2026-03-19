from __future__ import annotations

import csv
from pathlib import Path
import re
import subprocess
import tempfile

from docx import Document
from docx.enum.section import WD_ORIENTATION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt
from pdf2image import convert_from_path


STUDY_ROOT = Path(__file__).resolve().parents[1]
PACKAGE_DIR = STUDY_ROOT / "manuscript_package"


def set_document_defaults(document: Document, *, landscape: bool = False) -> None:
    section = document.sections[0]
    if landscape:
        section.orientation = WD_ORIENTATION.LANDSCAPE
        section.page_width, section.page_height = section.page_height, section.page_width
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)

    styles = document.styles
    normal = styles["Normal"]
    normal.font.name = "Times New Roman"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.15

    for style_name, size in [("Title", 16), ("Heading 1", 14), ("Heading 2", 12), ("Heading 3", 11), ("List Bullet", 11)]:
        style = styles[style_name]
        style.font.name = "Times New Roman"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
        style.font.size = Pt(size)
        if style_name != "List Bullet":
            style.font.bold = True


def finalize_paragraph(paragraph, *, alignment: WD_ALIGN_PARAGRAPH | None = None, first_line_indent: Inches | None = Inches(0.25)) -> None:
    if alignment is not None:
        paragraph.alignment = alignment
    elif paragraph.style.name == "Normal":
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    if paragraph.style.name == "Normal":
        paragraph.paragraph_format.first_line_indent = first_line_indent
    else:
        paragraph.paragraph_format.first_line_indent = None


def add_markdown_block(document: Document, line: str) -> None:
    stripped = line.strip()
    if not stripped:
        document.add_paragraph("")
        return
    stripped = stripped.replace("`", "")
    stripped = re.sub(r"\*\*(.*?)\*\*", r"\1", stripped)
    if stripped.startswith("# "):
        p = document.add_paragraph(stripped[2:], style="Title")
        finalize_paragraph(p, alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=None)
        return
    if stripped.startswith("## "):
        p = document.add_paragraph(stripped[3:], style="Heading 1")
        finalize_paragraph(p, alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=None)
        return
    if stripped.startswith("### "):
        p = document.add_paragraph(stripped[4:], style="Heading 2")
        finalize_paragraph(p, alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=None)
        return
    if stripped.startswith("- "):
        p = document.add_paragraph(stripped[2:], style="List Bullet")
        finalize_paragraph(p, alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=None)
        return
    if stripped.startswith("**Keywords:**"):
        p = document.add_paragraph()
        run = p.add_run("Keywords: ")
        run.bold = True
        p.add_run(stripped.split("**Keywords:**", 1)[1].strip())
        finalize_paragraph(p, alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=None)
        return
    p = document.add_paragraph(stripped)
    finalize_paragraph(p)


def markdown_to_docx(markdown_path: Path, output_path: Path) -> None:
    document = Document()
    set_document_defaults(document)
    for line in markdown_path.read_text(encoding="utf-8").splitlines():
        add_markdown_block(document, line)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def text_to_docx(text_path: Path, title: str, output_path: Path, bullets: bool = False) -> None:
    document = Document()
    set_document_defaults(document)
    title_p = document.add_paragraph(title, style="Title")
    finalize_paragraph(title_p, alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=None)
    for line in text_path.read_text(encoding="utf-8").splitlines():
        cleaned = line.strip()
        if not cleaned:
            continue
        if bullets:
            if cleaned.startswith("- "):
                cleaned = cleaned[2:]
            p = document.add_paragraph(cleaned, style="List Bullet")
            finalize_paragraph(p, alignment=WD_ALIGN_PARAGRAPH.LEFT, first_line_indent=None)
        else:
            p = document.add_paragraph(cleaned)
            finalize_paragraph(p)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def add_picture_block(document: Document, image_path: Path, *, width: Inches) -> None:
    paragraph = document.add_paragraph()
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    run.add_picture(str(image_path), width=width)
    finalize_paragraph(paragraph, alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=None)


def add_csv_table_block(document: Document, csv_path: Path) -> None:
    with csv_path.open("r", encoding="utf-8-sig", newline="") as handle:
        rows = list(csv.reader(handle))
    if not rows:
        return
    n_cols = max(len(row) for row in rows)
    table = document.add_table(rows=len(rows), cols=n_cols)
    table.style = "Table Grid"
    table.autofit = True
    for row_idx, row in enumerate(rows):
        for col_idx in range(n_cols):
            value = row[col_idx] if col_idx < len(row) else ""
            cell = table.cell(row_idx, col_idx)
            cell.text = value
            paragraph = cell.paragraphs[0]
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER if row_idx == 0 else WD_ALIGN_PARAGRAPH.LEFT
            paragraph.paragraph_format.first_line_indent = None
            if row_idx == 0:
                for run in paragraph.runs:
                    run.bold = True
    document.add_paragraph("")


def find_si_figure(figures_dir: Path, index: int) -> Path | None:
    for suffix in [".tiff", ".png", ".pdf"]:
        matches = sorted(figures_dir.glob(f"sifig{index:02d}_*{suffix}"))
        if matches:
            return matches[0]
    return None


def find_si_table(tables_dir: Path, index: int) -> Path | None:
    matches = sorted(tables_dir.glob(f"tableS{index}_*.xlsx"))
    return matches[0] if matches else None


def render_xlsx_pages(xlsx_path: Path, output_dir: Path) -> list[Path]:
    output_dir.mkdir(parents=True, exist_ok=True)
    pdf_dir = output_dir / "pdf"
    pdf_dir.mkdir(parents=True, exist_ok=True)
    subprocess.run(
        [
            "soffice",
            "--headless",
            "--convert-to",
            "pdf",
            "--outdir",
            str(pdf_dir),
            str(xlsx_path),
        ],
        check=True,
        stdout=subprocess.PIPE,
        stderr=subprocess.PIPE,
    )
    pdf_path = pdf_dir / f"{xlsx_path.stem}.pdf"
    images = convert_from_path(str(pdf_path), dpi=150)
    page_paths: list[Path] = []
    for page_index, image in enumerate(images, start=1):
        page_path = output_dir / f"{xlsx_path.stem}_p{page_index:02d}.jpg"
        rgb = image.convert("RGB")
        rgb.save(page_path, "JPEG", quality=82, optimize=True)
        page_paths.append(page_path)
    return page_paths


def supplementary_to_docx(markdown_path: Path, output_path: Path) -> None:
    document = Document()
    set_document_defaults(document, landscape=True)
    tables_dir = markdown_path.parent / "si_tables"
    figures_dir = markdown_path.parent / "si_figures"
    inserted_asset = False
    with tempfile.TemporaryDirectory(prefix="si_docx_assets_") as tmp_dir_name:
        tmp_dir = Path(tmp_dir_name)
        rendered_table_pages: dict[int, list[Path]] = {}

        for line in markdown_path.read_text(encoding="utf-8").splitlines():
            stripped = line.strip()
            figure_match = re.match(r"-\s*Figure S(\d+)\.", stripped)
            table_match = re.match(r"-\s*Table S(\d+)\.", stripped)

            if figure_match or table_match:
                if inserted_asset:
                    document.add_page_break()
                add_markdown_block(document, line)
                inserted_asset = True
                if figure_match:
                    figure_path = find_si_figure(figures_dir, int(figure_match.group(1)))
                    if figure_path is not None:
                        add_picture_block(document, figure_path, width=Inches(8.8))
                    else:
                        add_markdown_block(document, f"Missing asset: Figure S{figure_match.group(1)}")
                else:
                    table_index = int(table_match.group(1))
                    table_path = find_si_table(tables_dir, table_index)
                    if table_path is not None:
                        pages = rendered_table_pages.get(table_index)
                        if pages is None:
                            pages = render_xlsx_pages(table_path, tmp_dir / f"tableS{table_index}")
                            rendered_table_pages[table_index] = pages
                        for page_number, page_path in enumerate(pages, start=1):
                            if page_number > 1:
                                document.add_page_break()
                            add_picture_block(document, page_path, width=Inches(8.8))
                    else:
                        add_markdown_block(document, f"Missing asset: Table S{table_match.group(1)}")
                continue

            add_markdown_block(document, line)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(output_path)


def run_build(name: str, builder) -> str | None:
    try:
        builder()
    except PermissionError as exc:
        return f"{name}: {exc}"
    return None


def main() -> int:
    failures = []
    tasks = [
        ("manuscript.docx", lambda: markdown_to_docx(PACKAGE_DIR / "main_manuscript" / "manuscript.md", PACKAGE_DIR / "main_manuscript" / "manuscript.docx")),
        ("SI.docx", lambda: supplementary_to_docx(PACKAGE_DIR / "supplementary_information" / "SI.md", PACKAGE_DIR / "supplementary_information" / "SI.docx")),
        ("figure_captions.docx", lambda: markdown_to_docx(PACKAGE_DIR / "figures" / "figure_captions.md", PACKAGE_DIR / "figures" / "figure_captions.docx")),
        ("graphical_abstract_brief.docx", lambda: markdown_to_docx(PACKAGE_DIR / "figures" / "graphical_abstract_brief.md", PACKAGE_DIR / "figures" / "graphical_abstract_brief.docx")),
        ("table_captions.docx", lambda: markdown_to_docx(PACKAGE_DIR / "tables" / "table_captions.md", PACKAGE_DIR / "tables" / "table_captions.docx")),
        ("submission_notes.docx", lambda: markdown_to_docx(PACKAGE_DIR / "main_manuscript" / "submission_notes.md", PACKAGE_DIR / "main_manuscript" / "submission_notes.docx")),
        ("cover_letter_draft.docx", lambda: markdown_to_docx(PACKAGE_DIR / "main_manuscript" / "cover_letter_draft.md", PACKAGE_DIR / "main_manuscript" / "cover_letter_draft.docx")),
        ("water_research_submission_checklist.docx", lambda: markdown_to_docx(PACKAGE_DIR / "water_research_submission_checklist.md", PACKAGE_DIR / "water_research_submission_checklist.docx")),
        ("highlights.docx", lambda: text_to_docx(PACKAGE_DIR / "main_manuscript" / "highlights.txt", "Highlights", PACKAGE_DIR / "main_manuscript" / "highlights.docx", bullets=True)),
        ("keywords.docx", lambda: text_to_docx(PACKAGE_DIR / "main_manuscript" / "keywords.txt", "Keywords", PACKAGE_DIR / "main_manuscript" / "keywords.docx")),
    ]
    for name, builder in tasks:
        failure = run_build(name, builder)
        if failure is not None:
            failures.append(failure)
    if failures:
        for failure in failures:
            print(failure)
        return 1
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
